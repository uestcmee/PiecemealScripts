# -*-coding:gbk*-
from os import listdir
import docx
# from docx.enum.table import WD_TABLE_ALIGNMENT  # 这两个是表格对齐方式，暂时没用
# from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd


class excelToWord():
    def __init__(self):
        # 修改的表格位置
        self.INDEX = 13
        # 文件存放位置
        self.readfile_path = r'./0_example_doc/'
        self.excel_source = r'./1_excel_source/'
        # 运行程序
        self.write_doc()
        pass

    def read_excel(self,doc):  # 读取excel文件
        df=pd.read_excel(self.excel_source+doc+'.xlsx',header=None)  # 读取包括标题行&标题列
        return df

    def write_doc(self):
        xlslist = listdir(self.excel_source)  # excel_source文件夹中所有文件
        for myxls in xlslist:
            docname=myxls[:-5]  # 去掉后缀，获得文件名
            df = self.read_excel(docname)  # 读取excel文件
            newdocx = docx.Document(self.readfile_path+'example.docx')
            # newdocx.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table = newdocx.tables  # 获取doc中的所有表格
            for index,oTable in enumerate(table):
                if index!=self.INDEX:  # 只修改第14个表格，对应盈利能力
                    continue
                rows_num = len(oTable.rows)  # 表格行数
                columns_num = len(oTable.columns)  # 表格列数
                # 循环行列，修改表格值
                # 如果不希望改变首行首列，可以改变循环起始点
                for i in range(0,rows_num):  # 行循环
                    for j in range(0,columns_num):  # 列循环
                        cell=df.iloc[i,j]
                        if pd.isnull(cell)==True:  # 如果为空值，输出空值
                            oTable.cell(i,j).text=''
                        else:
                            oTable.cell(i,j).text=str(cell)# 写入cell
                        # 调整表格对齐方式
                        # oTable.cell(i,j).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            newdocx.save('./2_new_doc/' + str(docname)+'.docx') # 保存doc到新文件夹中
            print('save '+str(myxls))

if __name__ == '__main__':
    excelToWord()